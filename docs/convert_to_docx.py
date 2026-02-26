"""
NightFuel - Markdown to DOCX Converter (Beautiful Edition)
Usage: python docs/convert_to_docx.py
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ── Paths ──────────────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
INPUT_MD    = os.path.join(SCRIPT_DIR, "NightFuel-Project-Doc.md")
OUTPUT_DOCX = os.path.join(SCRIPT_DIR, "NightFuel-Project-Doc.docx")

# ── Palette ────────────────────────────────────────────────────────────────────
C_ORANGE      = "F97316"   # brand primary
C_ORANGE_DARK = "EA6C0A"   # hover
C_ORANGE_LIGHT= "FED7AA"   # orange-200
C_NAVY        = "0F172A"   # slate-900
C_DARK        = "1E293B"   # slate-800
C_MID         = "475569"   # slate-600
C_LIGHT       = "94A3B8"   # slate-400
C_BORDER      = "E2E8F0"   # slate-200
C_CODE_BG     = "1E293B"   # dark code bg
C_CODE_FG     = "E2E8F0"   # light code text
C_ROW_ALT     = "FFF7ED"   # orange-50 alt row
C_ROW_WHITE   = "FFFFFF"
C_HDR_BG      = "F97316"   # table header bg

def rgb(h): return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


# ── Low-level XML helpers ──────────────────────────────────────────────────────
def shade_para(para, fill: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)

def shade_cell(cell, fill: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=90, right=90):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in (("top",top),("bottom",bottom),("left",left),("right",right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)

def cell_border(cell, color="E2E8F0", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tbl  = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tbl.append(el)
    tcPr.append(tbl)

def para_border_left(para, color: str, sz="18", space="4"):
    """Thick coloured left border on a paragraph (accent bar)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    lft  = OxmlElement("w:left")
    lft.set(qn("w:val"),   "single")
    lft.set(qn("w:sz"),    sz)
    lft.set(qn("w:space"), space)
    lft.set(qn("w:color"), color)
    pBdr.append(lft)
    pPr.append(pBdr)

def para_border_bottom(para, color: str, sz="6", space="4"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:space"), space)
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def add_page_number(para):
    """Insert 'Page X of Y' field into a paragraph."""
    run = para.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(C_LIGHT)
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    para._p.append(fld)
    ins = OxmlElement("w:instrText"); ins.text = " PAGE "; ins.set(qn("xml:space"),"preserve")
    para._p.append(ins)
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    para._p.append(fld2)
    run2 = para.add_run(" of ")
    run2.font.size = Pt(9); run2.font.color.rgb = rgb(C_LIGHT)
    fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "begin")
    para._p.append(fld3)
    ins2 = OxmlElement("w:instrText"); ins2.text = " NUMPAGES "; ins2.set(qn("xml:space"),"preserve")
    para._p.append(ins2)
    fld4 = OxmlElement("w:fldChar"); fld4.set(qn("w:fldCharType"), "end")
    para._p.append(fld4)


# ── Inline markup ──────────────────────────────────────────────────────────────
def apply_inline(container, text: str, base_size: Pt = Pt(11), base_color: str = C_DARK):
    """Parse **bold**, *italic*, `code` markers into styled runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    for part in pattern.split(text):
        if not part:
            continue
        run = container.add_run()
        run.font.name = "Calibri"
        run.font.size = base_size
        run.font.color.rgb = rgb(base_color)
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = rgb("D97706")   # amber-600
            run.bold = False
        elif part.startswith("*") and part.endswith("*"):
            run.text   = part[1:-1]
            run.italic = True
        else:
            run.text = part


# ── Document setup ─────────────────────────────────────────────────────────────
def make_doc() -> Document:
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin   = Cm(2.8)
        sec.right_margin  = Cm(2.8)

    # Normal style
    n = doc.styles["Normal"]
    n.font.name  = "Calibri"
    n.font.size  = Pt(11)
    n.font.color.rgb = rgb(C_DARK)
    n.paragraph_format.space_after  = Pt(6)
    n.paragraph_format.line_spacing = Pt(16)

    # Heading styles
    cfg = {
        "Heading 1": (Pt(24), True,  C_ORANGE,  Pt(22), Pt(6)),
        "Heading 2": (Pt(16), True,  C_NAVY,    Pt(16), Pt(4)),
        "Heading 3": (Pt(13), True,  C_DARK,    Pt(12), Pt(3)),
        "Heading 4": (Pt(11), True,  C_MID,     Pt(10), Pt(2)),
    }
    for name, (sz, bold, col, sb, sa) in cfg.items():
        s = doc.styles[name]
        s.font.name  = "Calibri"
        s.font.size  = sz
        s.font.bold  = bold
        s.font.color.rgb = rgb(col)
        s.paragraph_format.space_before = sb
        s.paragraph_format.space_after  = sa

    return doc


# ── Cover page ─────────────────────────────────────────────────────────────────
def clean_spacer(doc: Document, height_pt: float = 14):
    """Plain Normal-style spacer paragraph — no inherited borders or styles."""
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(" ")
    r.font.size = Pt(height_pt)
    return p

def remove_table_borders(table):
    """Set all table borders to nil (truly invisible — works in WPS & Word)."""
    tbl_el = table._tbl
    tblPr  = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl_el.insert(0, tblPr)
    # Remove any existing tblBorders element first
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    bdr = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "nil")   # "nil" = truly no border (WPS + Word)
        bdr.append(el)
    tblPr.append(bdr)
    # Also nil out every cell border
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tcBdr = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right","insideH","insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "nil")
                tcBdr.append(el)
            tcPr.append(tcBdr)

def add_cover(doc: Document):
    # Zero top margin for cover section so orange bar bleeds to edge
    cover_section = doc.sections[0]
    cover_section.top_margin    = Cm(0)
    cover_section.bottom_margin = Cm(0)
    cover_section.left_margin   = Cm(2.8)
    cover_section.right_margin  = Cm(2.8)

    # ── Top orange bar ────────────────────────────────────────────────────────
    bar = doc.add_paragraph(style="Normal")
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after  = Pt(0)
    shade_para(bar, C_ORANGE)
    run = bar.add_run(" " * 160)
    run.font.size = Pt(18)        # taller bar

    clean_spacer(doc, 28)
    clean_spacer(doc, 28)
    clean_spacer(doc, 28)

    # ── "NightFuel" title ─────────────────────────────────────────────────────
    p_title = doc.add_paragraph(style="Normal")
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(0)
    r = p_title.add_run("NightFuel")
    r.font.name  = "Calibri"
    r.font.size  = Pt(52)
    r.font.bold  = True
    r.font.color.rgb = rgb(C_ORANGE)

    clean_spacer(doc, 8)

    # ── Tagline ───────────────────────────────────────────────────────────────
    p_tag = doc.add_paragraph(style="Normal")
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tag.paragraph_format.space_before = Pt(0)
    p_tag.paragraph_format.space_after  = Pt(0)
    r2 = p_tag.add_run("Chrono-Nutrition & Fitness Platform for Shift Workers")
    r2.font.name  = "Calibri"
    r2.font.size  = Pt(16)
    r2.font.color.rgb = rgb(C_MID)
    r2.italic = True

    clean_spacer(doc, 14)

    # ── Orange divider ────────────────────────────────────────────────────────
    div = doc.add_paragraph(style="Normal")
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_before = Pt(0)
    div.paragraph_format.space_after  = Pt(0)
    para_border_bottom(div, C_ORANGE, sz="12", space="4")
    div.add_run(" ")

    clean_spacer(doc, 14)
    clean_spacer(doc, 14)

    # ── Meta table ────────────────────────────────────────────────────────────
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pairs = [
        ("Version",      "2.0"),
        ("Date",         "February 2026"),
        ("Stack",        "TypeScript 5.5  +  Python 3.12  +  Next.js 15"),
        ("Architecture", "14 Microservices  +  Turborepo Monorepo"),
    ]
    for r_idx, (label, value) in enumerate(pairs):
        row      = meta_table.rows[r_idx]
        lbl_cell = row.cells[0]
        val_cell = row.cells[1]

        lbl_cell.text = ""
        lp = lbl_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.name  = "Calibri"
        lr.font.size  = Pt(10.5)
        lr.font.color.rgb = rgb(C_ORANGE)

        val_cell.text = ""
        vp = val_cell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.name  = "Calibri"
        vr.font.size  = Pt(10.5)
        vr.font.color.rgb = rgb(C_DARK)

        for cell in (lbl_cell, val_cell):
            shade_cell(cell, C_ROW_WHITE)
            set_cell_margins(cell, top=50, bottom=50, left=100, right=100)

    remove_table_borders(meta_table)

    clean_spacer(doc, 28)
    clean_spacer(doc, 28)
    clean_spacer(doc, 28)
    clean_spacer(doc, 28)
    clean_spacer(doc, 28)
    clean_spacer(doc, 28)

    # ── Bottom navy bar ───────────────────────────────────────────────────────
    bot = doc.add_paragraph(style="Normal")
    bot.paragraph_format.space_before = Pt(0)
    bot.paragraph_format.space_after  = Pt(0)
    shade_para(bot, C_NAVY)
    br = bot.add_run("   NightFuel  |  Project Design Document v2.0  |  Confidential")
    br.font.name  = "Calibri"
    br.font.size  = Pt(9)
    br.font.color.rgb = rgb(C_LIGHT)

    # Page break — restore normal margins for all subsequent pages
    doc.add_page_break()
    for sec in doc.sections[1:]:
        sec.top_margin    = Cm(2.2)
        sec.bottom_margin = Cm(2.2)


# ── Header & Footer ────────────────────────────────────────────────────────────
def setup_header_footer(doc: Document):
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("NightFuel  |  Project Design Document v2.0")
        run.font.name  = "Calibri"
        run.font.size  = Pt(8.5)
        run.font.color.rgb = rgb(C_LIGHT)
        para_border_bottom(p, C_BORDER, sz="4", space="4")

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(fp)
        para_border_bottom(fp, C_BORDER, sz="4", space="4")


# ── Heading renderer ───────────────────────────────────────────────────────────
def add_heading(doc: Document, raw: str, level: int):
    clean = re.sub(r"^#+\s*", "", raw).strip()

    if level == 1:
        # Large orange heading with thick left accent bar
        para = doc.add_heading("", level=1)
        para.paragraph_format.space_before = Pt(24)
        para.paragraph_format.space_after  = Pt(8)
        para.paragraph_format.left_indent  = Inches(0.2)
        r = para.add_run(clean)
        r.font.name  = "Calibri"
        r.font.size  = Pt(22)
        r.font.bold  = True
        r.font.color.rgb = rgb(C_ORANGE)
        para_border_left(para, C_ORANGE, sz="24", space="8")
        para_border_bottom(para, C_ORANGE_LIGHT, sz="4", space="6")

    elif level == 2:
        para = doc.add_heading("", level=2)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after  = Pt(6)
        # Small orange square prefix
        r0 = para.add_run("  ")
        r0.font.size = Pt(6)
        r = para.add_run(clean)
        r.font.name  = "Calibri"
        r.font.size  = Pt(15)
        r.font.bold  = True
        r.font.color.rgb = rgb(C_NAVY)
        para_border_left(para, C_ORANGE, sz="12", space="6")

    elif level == 3:
        para = doc.add_heading("", level=3)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after  = Pt(4)
        para.paragraph_format.left_indent  = Inches(0.1)
        r = para.add_run(clean)
        r.font.name  = "Calibri"
        r.font.size  = Pt(12)
        r.font.bold  = True
        r.font.color.rgb = rgb(C_DARK)

    else:
        para = doc.add_heading("", level=4)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(2)
        r = para.add_run(clean)
        r.font.name  = "Calibri"
        r.font.size  = Pt(11)
        r.font.bold  = True
        r.font.italic = True
        r.font.color.rgb = rgb(C_MID)


# ── Code block ─────────────────────────────────────────────────────────────────
def add_code_block(doc: Document, lines: list, lang: str = ""):
    # Outer paragraph as dark container
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(8)
    para.paragraph_format.left_indent  = Inches(0.15)
    para.paragraph_format.right_indent = Inches(0.15)

    shade_para(para, C_CODE_BG)
    para_border_left(para, C_ORANGE, sz="20", space="6")

    # Language label (if any)
    if lang:
        lr = para.add_run(f"{lang}\n")
        lr.font.name  = "Consolas"
        lr.font.size  = Pt(7.5)
        lr.font.color.rgb = rgb(C_ORANGE)
        lr.bold = True

    code_text = "\n".join(lines)
    cr = para.add_run(code_text)
    cr.font.name  = "Consolas"
    cr.font.size  = Pt(9)
    cr.font.color.rgb = rgb(C_CODE_FG)


# ── Table ──────────────────────────────────────────────────────────────────────
def add_table(doc: Document, header_row: list, data_rows: list):
    n_cols = max(len(header_row), 1)
    table  = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove default Table Grid borders (we add custom ones)
    tbl_el = table._tbl
    tblPr  = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl_el.insert(0, tblPr)
    bdr_el = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), C_BORDER)
        bdr_el.append(el)
    tblPr.append(bdr_el)

    # Header row — dark navy background
    for i, col_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(col_text.strip())
        r.bold = True
        r.font.name  = "Calibri"
        r.font.size  = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, C_NAVY)
        set_cell_margins(cell)
        cell_border(cell, C_NAVY)

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        row = table.add_row()
        fill = C_ROW_ALT if row_idx % 2 == 0 else C_ROW_WHITE
        for i in range(n_cols):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            if i < len(row_data):
                apply_inline(p, row_data[i].strip(), Pt(10), C_DARK)
            shade_cell(cell, fill)
            set_cell_margins(cell)
            cell_border(cell, C_BORDER)

    doc.add_paragraph().paragraph_format.space_before = Pt(2)


# ── Bullet & numbered list ─────────────────────────────────────────────────────
def add_bullet(doc: Document, text: str, level: int = 0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent  = Inches(0.3 + 0.2 * level)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    apply_inline(para, text.strip(), Pt(11), C_DARK)
    # Colour the bullet character orange by injecting rPr on numPr run
    for r in para.runs:
        if not r.text.strip():
            r.font.color.rgb = rgb(C_ORANGE)

def add_numbered(doc: Document, text: str, level: int = 0):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.left_indent  = Inches(0.3 + 0.2 * level)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    apply_inline(para, text.strip(), Pt(11), C_DARK)


# ── Body paragraph ─────────────────────────────────────────────────────────────
def add_body(doc: Document, text: str):
    if not text.strip():
        return
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after  = Pt(5)
    para.paragraph_format.line_spacing = Pt(16)
    apply_inline(para, text, Pt(11), C_DARK)


# ── Horizontal rule ────────────────────────────────────────────────────────────
def add_hr(doc: Document):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(8)
    para_border_bottom(para, C_ORANGE_LIGHT, sz="6", space="6")


# ── Main converter ─────────────────────────────────────────────────────────────
def convert(md_path: str, docx_path: str):
    doc = make_doc()
    add_cover(doc)
    setup_header_footer(doc)

    lines = open(md_path, encoding="utf-8").readlines()

    in_code    = False
    code_buf   = []
    code_lang  = ""
    table_hdr  = []
    table_rows = []
    in_table   = False

    # Skip the document's own H1 title (it's on the cover page already)
    start_idx = 0
    for idx, l in enumerate(lines):
        if l.startswith("# "):
            start_idx = idx + 1
            break

    i = start_idx
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # ── Fenced code ──────────────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code:
                in_code   = True
                code_lang = line[3:].strip()
                code_buf  = []
            else:
                add_code_block(doc, code_buf, code_lang)
                in_code   = False
                code_buf  = []
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── HR ────────────────────────────────────────────────────────────────
        if re.match(r"^---+\s*$", line):
            add_hr(doc)
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            add_heading(doc, line, level)
            i += 1
            continue

        # ── Tables ───────────────────────────────────────────────────────────
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(re.match(r"^[-: ]+$", c) for c in cells if c):
                i += 1
                continue
            if not in_table:
                in_table   = True
                table_hdr  = cells
                table_rows = []
            else:
                table_rows.append(cells)
            next_l = lines[i + 1].rstrip("\n") if i + 1 < len(lines) else ""
            if not next_l.startswith("|"):
                add_table(doc, table_hdr, table_rows)
                in_table = False
            i += 1
            continue
        else:
            if in_table:
                add_table(doc, table_hdr, table_rows)
                in_table = False

        # ── Bullet ───────────────────────────────────────────────────────────
        m = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if m:
            add_bullet(doc, m.group(2), len(m.group(1)) // 2)
            i += 1
            continue

        # ── Numbered ─────────────────────────────────────────────────────────
        m = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if m:
            add_numbered(doc, m.group(2), len(m.group(1)) // 2)
            i += 1
            continue

        # ── Blank ────────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Body ─────────────────────────────────────────────────────────────
        add_body(doc, line)
        i += 1

    doc.save(docx_path)
    print(f"[OK] Saved: {docx_path}")


if __name__ == "__main__":
    convert(INPUT_MD, OUTPUT_DOCX)
